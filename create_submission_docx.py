import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Zyro AI/ML Internship")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Week 1: AI/ML Environment Setup & Verification Submission")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    meta_info = [
        ("Candidate Name / GitHub Username", "sohaib61595"),
        ("GitHub Profile Link", "https://github.com/sohaib61595"),
        ("GitHub Repository Link", "https://github.com/sohaib61595/zyro-aiml-internship"),
        ("Submission Milestone", "Week 01 - AI/ML Environment & Repository Setup")
    ]
    
    for row_idx, (label, val) in enumerate(meta_info):
        cell_0 = table.cell(row_idx, 0)
        cell_1 = table.cell(row_idx, 1)
        
        cell_0.width = Inches(2.5)
        cell_1.width = Inches(4.3)
        
        set_cell_background(cell_0, "EBF1F5")
        set_cell_background(cell_1, "FAFAFA")
        
        p0 = cell_0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10.5)
        
        p1 = cell_1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10.5)
        if "http" in val:
            r1.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
            r1.font.underline = True
            
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    
    # Section Header Helper
    def add_section_header(num, text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h_run = h.add_run(f"{num}. {text}")
        h_run.font.size = Pt(14)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        
    def add_image_box(img_path, caption="", width=Inches(5.8)):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(img_path, width=width)
            
            if caption:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(12)
                c_run = cp.add_run(caption)
                c_run.font.size = Pt(9.5)
                c_run.font.italic = True
                c_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            p = doc.add_paragraph()
            p.add_run(f"[Image not found at {img_path}]").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # 1. GitHub Profile Link
    add_section_header("1", "GitHub Profile Link")
    p = doc.add_paragraph()
    r = p.add_run("🔗 GitHub Profile: ")
    r.font.bold = True
    r2 = p.add_run("https://github.com/sohaib61595")
    r2.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    r2.font.underline = True

    # 2. GitHub Repository Link
    add_section_header("2", "GitHub Repository Link")
    p = doc.add_paragraph()
    r = p.add_run("📦 GitHub Repository: ")
    r.font.bold = True
    r2 = p.add_run("https://github.com/sohaib61595/zyro-aiml-internship")
    r2.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    r2.font.underline = True

    # 3. Screenshot of Python Version
    add_section_header("3", "Screenshot of Python Version")
    add_image_box("week-01/Screenshots/Python_Version.png", "Figure 1: Python 3.13+ Installation and Version Verification", width=Inches(5.6))

    # 4. Screenshot of Git Version
    add_section_header("4", "Screenshot of Git Version")
    add_image_box("week-01/Screenshots/git_version.png", "Figure 2: Git Version Verification and Setup", width=Inches(5.6))

    # 5. Screenshot of Virtual Environment
    add_section_header("5", "Screenshot of Virtual Environment (.venv)")
    add_image_box("week-01/Screenshots/virtual_environment.png", "Figure 3: Virtual Environment (.venv) Activation & Installed Packages", width=Inches(5.8))

    # 6. Screenshot of Successful ML Test
    add_section_header("6", "Screenshot of Successful ML Test")
    add_image_box("week-01/Screenshots/successful_ml_test.png", "Figure 4: Successful Machine Learning Pipeline Execution (Iris Dataset - Random Forest)", width=Inches(5.8))

    # Optional Bonus Visualizations
    if os.path.exists("week-01/graphs/ml_confusion_matrix.png"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("Generated Evaluation Visualizations:").font.bold = True
        add_image_box("week-01/graphs/ml_confusion_matrix.png", "Figure 5: Random Forest Classifier Confusion Matrix Heatmap", width=Inches(4.5))

    # 7. Proof of Joining the Zyroo Community
    add_section_header("7", "Proof of Joining the Zyroo Community")
    add_image_box("week-01/Screenshots/community_join.png", "Figure 6: Proof of Joining the Zyroo WhatsApp Community", width=Inches(4.2))

    # 8. Proof of Joining the AI/ML Channel
    add_section_header("8", "Proof of Joining the AI/ML Channel")
    add_image_box("week-01/Screenshots/channel_join.png", "Figure 7: Proof of Joining the Zyroo AI/ML Channel", width=Inches(4.2))

    output_path = "Zyro_AIML_Internship_Week_1_Submission.docx"
    doc.save(output_path)
    print(f"Document successfully created: {output_path}")

    # Also save a copy inside week-01/
    doc.save("week-01/Zyro_AIML_Internship_Week_1_Submission.docx")
    print("Saved copy to week-01/Zyro_AIML_Internship_Week_1_Submission.docx")

if __name__ == "__main__":
    create_document()
